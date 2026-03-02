import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    # Style standardizations
    for run in heading.runs:
        run.font.name = 'Calibri'
        if level == 0:
            run.font.size = Pt(24)
            run.font.color.rgb = RGBColor(0, 51, 102) # Dark Blue
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif level == 1:
            run.font.color.rgb = RGBColor(0, 102, 204)

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    return p

def create_master_doc():
    doc = Document()
    
    # Title
    add_heading(doc, 'SPORTS: Master Project Architecture & Analysis Document', 0)
    
    doc.add_paragraph() # Spacer
    
    # Intro
    p = add_paragraph(doc, 'SPORTS (Student Progress Organizer and Record Track of Study) ', bold=True)
    p.add_run('is a comprehensive student support platform designed to act as a dynamic, living roadmap for career growth. This document serves as the Master Output of the complete system analysis, covering the architecture, data flow, component design, latency evaluation, and future roadmap.').font.name = 'Calibri'

    # Section 1
    add_heading(doc, '1. Project Requirements & Objective', 1)
    add_paragraph(doc, 'The primary objective of SPORTS is to track, assess, and guide a student\'s educational and professional development through data-driven insights and AI integration.')
    
    add_heading(doc, 'Key Modules Evaluated:', 2)
    doc.add_paragraph('The Portfolio Vault (Unified Identity): Acts as a central digital identity replacing static resumes. Features a verifiable locker for certifications.', style='List Bullet')
    doc.add_paragraph('Generate Future Path (AI-Powered Growth Engine): Generates specialized learning paths calculating the shortest route to a target job role using predictive mapping.', style='List Bullet')
    doc.add_paragraph('The Test Lab (Adaptive Assessment): Replaces static quizzes with an adaptive competency-based validation system identifying specific skill gaps.', style='List Bullet')
    doc.add_paragraph('Placement & Interview Hub (Predictive Intelligence): Matches industry job requirements against verified student skills and simulates company-specific interviews.', style='List Bullet')
    doc.add_paragraph('Strategic Analytics (Resume & Radar): Tailors resumes per job role automatically and predicts hackathons the student has a high likelihood of winning.', style='List Bullet')

    # Section 2
    add_heading(doc, '2. Project Structure Analysis', 1)
    add_paragraph(doc, 'The architecture utilizes a decoupled Next.js (React) frontend and a FastAPI (Python) backend.')
    
    add_heading(doc, 'Frontend (/frontend)', 2)
    doc.add_paragraph('App Router (/app): Contains the core functional pages (login, signup, onboarding, dashboard, opportunities).', style='List Bullet')
    doc.add_paragraph('Components (/components): Reusable, highly modular React elements (e.g., JobDetailDrawer, SkillSelector).', style='List Bullet')
    doc.add_paragraph('State Management: Relies efficiently on React Contexts and hooks.', style='List Bullet')
    
    add_heading(doc, 'Backend (/backend)', 2)
    doc.add_paragraph('FastAPI Core (main.py): The primary server entry point managing CORS and mounting domain-specific routers.', style='List Bullet')
    doc.add_paragraph('Routers (/routers): Vertically sliced APIs handling chat, roadmap, opportunities, prep, and resume.', style='List Bullet')
    doc.add_paragraph('Scrapers (/scrapers): Asynchronous python scripts dynamically pulling competitive programming and job data from external sources.', style='List Bullet')

    # Section 3
    add_heading(doc, '3. Technical, Database & API Flow', 1)
    add_heading(doc, 'Database Schema (Supabase / PostgreSQL)', 2)
    doc.add_paragraph('profiles: The centralized user spine linked to Supabase Auth. Utilizes jsonb scaling for flexible skills and resume data.', style='List Bullet')
    doc.add_paragraph('jobs & hackathons: Bulletin board tables operating independently with Row Level Security (RLS) enabled.', style='List Bullet')

    add_heading(doc, 'API Connectivity', 2)
    doc.add_paragraph('Frontend → Supabase: Client-side mutations for rapid perceived performance.', style='List Bullet')
    doc.add_paragraph('Frontend → FastAPI: Forwarding heavy compute tasks (test evaluations, roadmap generation).', style='List Bullet')
    doc.add_paragraph('FastAPI → Gemini AI: Acts as the cognitive engine for unstructured resume parsing and complex learning tree generation.', style='List Bullet')

    # Section 4
    add_heading(doc, '4. Current Working Ecosystem Review', 1)
    doc.add_paragraph('Stability: Next.js and FastAPI start cleanly with no strict crashing build errors.', style='List Bullet')
    doc.add_paragraph('Defensive Programming: Backend routers use try-except ImportError logic to ensure application uptime even if a modular python file is displaced.', style='List Bullet')
    doc.add_paragraph('Document Engine: The /api/parse-resume correctly intercepts .docx, decodes it via python-docx, and pipes it into Gemini.', style='List Bullet')

    # Section 5
    add_heading(doc, '5. Process & Workflow Analysis', 1)
    doc.add_paragraph('The Resume Ingestion: Works flawlessly on .docx. It decodes bytes to strings locally before executing an AI call.', style='List Bullet')
    doc.add_paragraph('Market Scraping: Fully abstracted into independent endpoint triggers (/api/scrape/*). This guarantees that a slow job-board website won\'t accidentally hang a student\'s login request.', style='List Bullet')
    doc.add_paragraph('Competency Assessments: UI state securely gates the user\'s progress until completion, preventing spoofed test completions being sent to the backend.', style='List Bullet')

    # Section 6
    add_heading(doc, '6. Visual Presentation & UI', 1)
    doc.add_paragraph('SPA Composability: The Dashboard leverages seamless nested components (Overview, Roadmap, Test Lab) eliminating full page reloads.', style='List Bullet')
    doc.add_paragraph('Deep Integration Modules: Heavy reliance on Drawers and Modals to surface deep data while keeping the user anchored on their primary view.', style='List Bullet')
    doc.add_paragraph('Visual Roadmap Canvas: Geometrically splits into "Retrospective" (past skills) and "Predictive" (future skills) creating an immersive, gaming-style skill-tree experience.', style='List Bullet')

    # Section 7
    add_heading(doc, '7. Identifying Bugs & Latency Traps', 1)
    add_heading(doc, 'Potential Bugs', 2)
    doc.add_paragraph('PDF Support Omission: The backend explicitly rejects .pdf files. Since most resumes are PDFs, this causes immediate UX friction.', style='List Bullet')
    doc.add_paragraph('Scraper Brittleness: If an external job board updates its HTML tags, BeautifulSoup scrapers may return None and crash the data insertion loops.', style='List Bullet')
    
    add_heading(doc, 'Latency Traps', 2)
    doc.add_paragraph('Synchronous LLM Calls: /api/roadmap/generate forces the user to wait synchronously for Gemini. This can cause the browser to look frozen.', style='List Bullet')

    # Section 8
    add_heading(doc, '8. Recommended UI Fallback States', 1)
    doc.add_paragraph('Roadmap Timeouts: "We hit a snag mapping your future! Our AI is currently recalibrating market trends. Try generating again in a few moments."', style='List Bullet')
    doc.add_paragraph('Job Scraper Failures: "The job boards are currently out of reach. Don\'t worry, your Compatibility Scores are safe. While we reconnect, why not brush up your skills in the Test Lab?"', style='List Bullet')
    doc.add_paragraph('Resume Crashes: "Our Resume Engine is overloaded right now. Please hit retry, or proceed with your default profile."', style='List Bullet')

    # Section 9
    add_heading(doc, '9. Runtime Logs Strategy', 1)
    doc.add_paragraph('Backend (Python): Shift to standard JSON logging. Every error log must output the user_uuid, latency_ms, and endpoint so we can track exactly which endpoint the LLM failed on.', style='List Bullet')
    doc.add_paragraph('Frontend (Next.js): Wrap the application in Sentry.ErrorBoundary to capture device stats, browser cache states, and React component stack traces for white-screen failures.', style='List Bullet')

    # Section 10
    add_heading(doc, '10. Future Enhancements & Extensibility', 1)
    doc.add_paragraph('Code Colosseum: Introduce an embedded web-IDE (like Monaco) allowing real-time technical test simulations instead of just multiple choice questions.', style='List Bullet')
    doc.add_paragraph('Alumni Connect Network: A mentorship engine mapping freshmen lacking a skill to seniors who aced the Test Lab for that exact skill.', style='List Bullet')
    doc.add_paragraph('Real-Time Interview Sandbox: Use WebRTC browser microphones and Gemini Vocals to conduct a physically spoken AI interview.', style='List Bullet')

    # Save logic
    output_path = os.path.abspath(os.path.join(os.path.dirname(__file__), 'SPORTS_Master_Project_Report.docx'))
    doc.save(output_path)
    print(f"Successfully generated {output_path}")

if __name__ == "__main__":
    create_master_doc()
